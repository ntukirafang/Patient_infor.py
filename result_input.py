from docx.shared import Cm # ..........cm
import pandas as pd
from docx.enum.table import WD_TABLE_ALIGNMENT

#kraken_combine = "/mnt/dell5820/PFI/report/report_BGI/2021-05-25-01-10-24_449_auto/JB21_003_0519/JB21_003_0519.report.combine_list.tsv"
VFDB = "/mnt/libobio_MGI/analysis_output_pfi/2021-05-25-01-10-24_449/Result/JB21_003_0519/2.Annotation/2.VFDB/DNA/VFDB.filter.xlsx"
CARD = "/mnt/libobio_MGI/analysis_output_pfi/2021-05-25-01-10-24_449/Result/JB21_003_0519/2.Annotation/4.CARD/DNA/CARD.filter.xlsx"
report_template = "/home/jimmy/ICTV_MGI_compare/Auto_run_script/PFI_ISO_Report_Template_20210714_mod_head_Jimmy_0716.docx"
report = "/home/jimmy/ICTV_MGI_compare/Auto_run_script/PFI_ISO_Report_linux_V2.docx"
with io.open(sys.argv[1], encoding ='utf-8') as fp:
    all_lines = fp.readlines()
with io.open(sys.argv[2], encoding ='utf-8') as df:
        all_lines2 = df.readlines()
print(all_lines)
print(all_lines[0].replace('\n','')[0])
print(all_lines[1].replace('\n', ''))
print(all_lines[1].replace('\n','').split('\t')[0])
print(len(all_lines))
env_path = []
input("pause_AA")

df_VFDB = pd.read_excel(VFDB)
print(df_VFDB)
df_CARD = pd.read_excel(CARD)
print(df_CARD)
df_user_ID = pd.read_excel('user_info.xlsx')
print(df_user_ID)
input("pause")




docx_list = [report_template]
for doc in docx_list:
    print(doc)

    document = Document(doc) # .........
    print(Document(doc))


    section = document.sections[1]
    header = section.header
    tables = header.tables
    table = tables[1]

    tables = document.tables #.........
    table = tables[4] #........

    for x in range(0,len(all_lines)):
        count = 1
        y = 1
        if all_lines[x].replace('\n', '') == '#Bacteria_top20':
            while y < 20:
                if all_lines[x + y].replace('\n', '')[0] != '#' and count <= 10:
                    print("count: " + str(count))
                    print("y: " + str(y))
                    print(len(all_lines[x + y].split("\t")))
                    if len(all_lines[x + y].split("\t")) == 3:
                        table.cell(count, 1).paragraphs[0].add_run(all_lines[x + y].split("\t")[0]).italic = True
                        table.cell(count, 2).paragraphs[0].text = all_lines[x + y].split("\t")[1]
                        table.cell(count, 3).paragraphs[0].text = str(
                            round(float(all_lines[x + y].split("\t")[2]), 1))
                        count = count + 1
                        y = y + 1
#                    elif int(all_lines[x + y].split("\t")[1]) / int(all_lines[x + y].split("\t")[3]) >= 10:
#                        table.cell(count, 1).paragraphs[0].add_run(all_lines[x + y].split("\t")[0]).italic = True
#                        table.cell(count, 2).paragraphs[0].text = all_lines[x + y].split("\t")[1]
#                        table.cell(count, 3).paragraphs[0].text = str(
#                            round(float(all_lines[x + y].split("\t")[2]) * 100, 1))
#                        count = count + 1
#                        y = y + 1
                    else:
                        env_path.append(["Bacteria", all_lines[x + y].split("\t")[0], all_lines[x + y].split("\t")[3]])
                        y = y + 1
                else:
                    break

    print('#Fungi_top20')
    table = tables[6]  # ........
    for x in range(0,len(all_lines)):
        count = 1
        y = 1
        if all_lines[x].replace('\n','') == '#Fungi_top20':
            while y < 20:
                if all_lines[x + y].replace('\n', '')[0] != '#' and count <= 10:
                    print("count: " + str(count))
                    print("y: " + str(y))
                    print(all_lines[x + y].replace('\n', ''))
                    if len(all_lines[x + y].split("\t")) == 3:
                        table.cell(count, 1).paragraphs[0].add_run(all_lines[x + y].split("\t")[0]).italic = True
                        table.cell(count, 2).paragraphs[0].text = all_lines[x + y].split("\t")[1]
                        table.cell(count, 3).paragraphs[0].text = str(
                            round(float(all_lines[x + y].split("\t")[2]), 1))
                        count = count + 1
                        y = y + 1                   
#                    elif int(all_lines[x + y].split("\t")[1]) / int(all_lines[x + y].split("\t")[3]) >= 10 :
#                        table.cell(count, 1).paragraphs[0].add_run(all_lines[x + y].split("\t")[0]).italic = True
#                        table.cell(count, 2).paragraphs[0].text = all_lines[x + y].split("\t")[1]
#                        table.cell(count, 3).paragraphs[0].text = str(
#                            round(float(all_lines[x + y].split("\t")[2]) * 100, 1))
#                        count = count + 1
#                        y = y + 1
                    else:
                        env_path.append(["Fungi", all_lines[x + y].split("\t")[0], all_lines[x + y].split("\t")[3]])
                        y = y + 1
                else:
                    break
    print('#Virus_top20_DNA')
    res = requests.get("https://www.genome.jp/virushostdb/index/virus/all" )
    html = BeautifulSoup(res.text, 'html.parser')
    tables1 = html.findAll("table",{"border":"1"})
    table1 = tables1[0]
    tds = table1.findAll("td")
    total_index = tds[0].text
    Index = int(total_index.split(" ")[1])
    virus_string = ""
    virus_list1 = open('/home/jimmy/ICTV_MGI_compare/Auto_run_script/virus_list.txt', "w", encoding ='utf-8')
    for i in range(Index):
        virus_name = tds[i].text
        virus_list_file = virus_name.split(" (")[0].split()
        virus_string = " ".join(virus_list_file).lstrip()
        virus_list_name = virus_string.replace(" ","\t",1)
        virus_list1.write(virus_list_name + "\n")
    virus_list1.close()
    virus_list2 = open('/home/jimmy/ICTV_MGI_compare/Auto_run_script/virus_list.txt', "r", encoding ='utf-8')
    virus_ref_list = virus_list2.readlines()
    virus_list2.close()
    
    table = tables[8]# ........
    DNA_Status = open('/home/jimmy/ICTV_MGI_compare/Auto_run_script/2020_MGI_DNA.txt', "r")
    DNA_list = DNA_Status.readlines()
    DNA_Status.close()
    pos = all_lines.index('#Viruses_top20\n')
    end = all_lines.index('#end\n')
    virus_name = all_lines[int(pos)+ 1:int(end)]
    count = 0
    unclassfial_virus_DNA = open("/home/jimmy/ICTV_MGI_compare/Auto_run_script/unc_virus_DNA.txt",'w')
    virus_DNA = open("/home/jimmy/ICTV_MGI_compare/Auto_run_script/c_ICTV_virus_DNA.txt",'w')
    virus_DNA_1 = open("/home/jimmy/ICTV_MGI_compare/Auto_run_script/c_Virus_Db_virus_DNA.txt",'w') 
    class_ICTV_virus = []
    class_VirusDb_virus = []
    for virus in virus_name:
        if len(virus.split("\t")) == 3 and count <= 9:
            count = count +1
            table.cell(count, 1).paragraphs[0].add_run(virus.split("\t")[0]).italic = True
            table.cell(count, 2).paragraphs[0].text = virus.split("\t")[1]
            table.cell(count, 3).paragraphs[0].text = str(
                    round(float(virus.split("\t")[2]), 1))
            print(virus.split('\t')[0])
            print(count)
            input()
            for item1 in DNA_list:
                if virus.split('\t')[0] == item1.split('\t')[0]:
                    table.cell(count, 1).paragraphs[0].add_run("*")
                    print(virus.split('\t')[0])
                    print(count)
                    input()
                    print(virus.split("\t"), file = virus_DNA)
                    class_ICTV_virus.append(virus)
            for ref_list in virus_ref_list:
                if virus.split('\t')[0] == ref_list.split("\t")[1].strip():
                    taxid = ref_list.split("\t")[0]
                    taxname = ref_list.split("\t")[1]
                    res = requests.get("https://www.genome.jp/virushostdb/" + str(taxid))
                    html = BeautifulSoup(res.text, 'html.parser')
                    tables2 = html.findAll("table",{"border":"1"})
                    table2 = tables2[0]
                    tds = table2.findAll("td")
                    td = tds[2].text
                    print(td)
                    index_ID = td.find("NA")
                    DNA_or_RNA_type = td[index_ID - 1]
                    if DNA_or_RNA_type == 'R':
                        print(taxname, taxid, DNA_or_RNA_type,'RNA')
                    elif DNA_or_RNA_type == 'D':
                        print(taxname, taxid, DNA_or_RNA_type,'DNA')
                        table.cell(count, 1).paragraphs[0].add_run("#")
                        print(virus.split("\t"), file = virus_DNA_1)
                        class_VirusDb_virus.append(virus)
                else:
                    continue
        elif len(virus.split("\t")) == 5 and count <= 9:
            env_path.append(["Virus", virus.split("\t")[0], virus.split("\t")[1]])
        else:
            break
    for virus in virus_name[0:10]:
        if virus not in class_ICTV_virus and virus not in class_VirusDb_virus :
            print(virus.split("\t"), file = unclassfial_virus_DNA)
    unclassfial_virus_DNA.close()
    virus_DNA.close()
    virus_DNA_1.close()
    
    print('#Virus_top20_RNA')
    table = tables[10]
    RNA_Status = open('/home/jimmy/ICTV_MGI_compare/Auto_run_script/2020_MGI_RNA.txt', "r")
    RNA_list = RNA_Status.readlines()
    RNA_Status.close()
    pos = all_lines2.index('#Viruses_top20\n')
    end = all_lines2.index('#end\n')
    virus_name = all_lines2[int(pos)+ 1:int(end)]
    count = 0
    unclassfial_virus_RNA = open("/home/jimmy/ICTV_MGI_compare/Auto_run_script/unc_virus_RNA.txt",'w')
    virus_ICTV_RNA = open("/home/jimmy/ICTV_MGI_compare/Auto_run_script/c_ICTV_virus_RNA.txt",'w')
    virus_Virus_Db_RNA = open("/home/jimmy/ICTV_MGI_compare/Auto_run_script/c_Virus_Db_virus_RNA.txt",'w')
    class_ICTV_virus_RNA = []
    class_VirusDb_virus_RNA = []
    for virus in virus_name:
        if len(virus.split("\t")) == 3 and count <= 9:
            count = count +1
            table.cell(count, 1).paragraphs[0].add_run(virus.split("\t")[0]).italic = True
            table.cell(count, 2).paragraphs[0].text = virus.split("\t")[1]
            table.cell(count, 3).paragraphs[0].text = str(
                    round(float(virus.split("\t")[2]), 1))
            print(virus.split('\t')[0])
            print(count)
            input()
            for item2 in RNA_list:
                if virus.split('\t')[0] == item2.split('\t')[0]:
                    table.cell(count, 1).paragraphs[0].add_run("*")
                    print(virus.split('\t')[0])
                    print(count)
                    input()
                    print(virus.split("\t"), file = virus_ICTV_RNA)
                    class_ICTV_virus.append(virus)
            for ref_list in virus_ref_list:
                if virus.split('\t')[0] == ref_list.split("\t")[1].strip():
                    taxid = ref_list.split("\t")[0]
                    taxname = ref_list.split("\t")[1]
                    res = requests.get("https://www.genome.jp/virushostdb/" + str(taxid))
                    html = BeautifulSoup(res.text, 'html.parser')
                    tables2 = html.findAll("table",{"border":"1"})
                    table2 = tables2[0]
                    tds = table2.findAll("td")
                    td = tds[2].text
                    print(td)
                    index_ID = td.find("NA")
                    DNA_or_RNA_type = td[index_ID - 1]
                    if DNA_or_RNA_type == 'D':
                        print(taxname, taxid, DNA_or_RNA_type,'DNA')
                    elif DNA_or_RNA_type == 'R':
                        print(taxname, taxid, DNA_or_RNA_type,'RNA')
                        table.cell(count, 1).paragraphs[0].add_run("#")
                        print(virus.split("\t"), file = virus_Virus_Db_RNA)
                        class_VirusDb_virus.append(virus)
                else:
                    continue
        elif len(virus.split("\t")) == 5 and count <= 9:
            env_path.append(["Virus", virus.split("\t")[0], virus.split("\t")[1]])
        else:
            break
    for virus in virus_name[0:10]:
        if virus not in class_ICTV_virus and virus not in class_VirusDb_virus :
            print(virus.split("\t"), file = unclassfial_virus_RNA)
            unclassfial_virus_RNA.close()
            virus_ICTV_RNA.close()
            virus_Virus_Db_RNA.close()
#    table = tables[12]  # ........
#    if len(df_VFDB) <= 10:
#        for x in range(0, len(df_VFDB)):
#            for y in range(0,6):
#                table.cell(x + 1, y).paragraphs[0].add_run(str(df_VFDB.iloc[x,y]))
#    else:
#        for x in range(0, 10):
#            for y in range(0, 6):
#                table.cell(x + 1, y).paragraphs[0].add_run(str(df_VFDB.iloc[x,y]))
#                print("x,y: "+str(x)+','+str(y))
#                print(df_VFDB.iloc[x,y])
#
#    table = tables[14]  # ........
#
#    x=0
#    while x < len(df_CARD) and x <= 10:
#        for y in range(0, 8):
#          if y!= 4:
#            table.cell(x + 1, y).paragraphs[0].add_run(str(df_CARD.iloc[x, y]))
#            print("x,y: " + str(x) + ',' + str(y))
#            print(df_CARD.iloc[x, y])
#          else:
#            table.cell(x + 1, y).paragraphs[0].add_run(str(df_CARD.iloc[x, y])).italic = True
#            print("x,y: " + str(x) + ',' + str(y))
#            print(df_CARD.iloc[x, y])
#        x = x + 1
#        #input("pause")

    print(env_path)
    print(len(env_path))
    table = tables[18]
    x = 0
    while x < len(env_path) and x < 20:
       table.cell(x + 3, 2).paragraphs[0].add_run(str(env_path[x][0]))
       table.cell(x + 3, 3).paragraphs[0].add_run(str(env_path[x][1])).italic = True
       table.cell(x + 3, 5).paragraphs[0].add_run(str(env_path[x][2]))
       print("x: " + str(x))
       print(str(env_path[x][0]))
       print(str(env_path[x][1]))
       print(str(env_path[x][2]))
       x = x + 1
    document.save(report)
