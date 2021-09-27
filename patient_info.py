#!/usr/bin/env python
# -*- coding: utf-8 -*-
import io
import os
import sys
import requests
import docx
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches # ...........
from docx.shared import Cm # ..........cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
doc = docx.Document('/home/jimmy/report/patient_info/PFI_ISO_Report_Template_20210806_jimmy.docx')
result_file2 = pd.read_excel('/home/jimmy/report/patient_info/patient_info_list.xlsx')
result_file1 = result_file2.replace(to_replace = np.nan, 
                 value = "N/A")
list_of_sample_name = result_file1['檢體名稱'].tolist()
this_time_sample_name = input("Please type JB21_XXX :")
print(this_time_sample_name)
Index_of_sample = int(list_of_sample_name.index(this_time_sample_name))
print(type(Index_of_sample))
print(Index_of_sample)
list_of_name = result_file1['姓名'].tolist()
list_of_sex = result_file1['性別'].tolist()
print(list_of_sex)
list_of_sample_type = result_file1['檢體類型'].tolist()
list_of_patient_number = result_file1['病歷編號'].tolist()
list_of_test_time = result_file1['採集日期'].tolist()
list_of_receive_time = result_file1['收檢日期'].tolist()
list_of_report_time = result_file1['報告日期'].tolist()
list_of_birthday = result_file1['生日'].tolist()
list_of_ID = result_file1['身分證字號'].tolist()
list_of_hospital = result_file1['委託機構'].tolist()
list_of_doctor = result_file1['醫師'].tolist()
list_of_contact = result_file1['聯絡人'].tolist()
list_of_phone = result_file1['電話'].tolist()
list_of_mail = result_file1['信箱'].tolist()
section = doc.sections[1]
header = section.header
head_tables = header.tables
print(list_of_name[Index_of_sample])
patient_name = list_of_name[Index_of_sample]
print(patient_name)
patient_birthday = list_of_birthday[Index_of_sample]
print(patient_birthday)
patient_sex = list_of_sex[Index_of_sample]
print(patient_sex)
patient_report_ID = list_of_sample_name[Index_of_sample]
print(patient_report_ID)
patient_hospital = list_of_hospital[Index_of_sample]
print(patient_hospital)
patient_ID = list_of_ID[Index_of_sample]
print(patient_ID)
patient_sample_type = list_of_sample_type[Index_of_sample]
print(patient_sample_type)
patient_number = list_of_patient_number[Index_of_sample]
print(patient_number)
patient_doctor = list_of_doctor[Index_of_sample]
print(patient_doctor)
patient_test_time = list_of_test_time[Index_of_sample]
print(patient_test_time)
patient_contact = list_of_contact[Index_of_sample]
print(patient_contact)
patient_receive_time = list_of_receive_time[Index_of_sample]
print(patient_receive_time)
doctor_phone = list_of_phone[Index_of_sample]
print(list_of_phone)
print(doctor_phone)
patient_report_time = list_of_report_time[Index_of_sample]
print(patient_report_time)
doctor_email = list_of_mail[Index_of_sample]
print(doctor_email)
run = head_tables[1].cell(0,3).paragraphs[0].add_run(str(patient_name))
run.font.size = Pt(9)
run.font.bold = True
#header_姓名
run = head_tables[1].cell(0,5).paragraphs[0].add_run(str(patient_birthday))
run.font.name = 'Arial'
run.font.size = Pt(9)
run.font.bold = True
#header_生日
run = head_tables[1].cell(1,3).paragraphs[0].add_run(str(patient_sex))
#run.font.name = 'Microsoft JhengHei'
run.font.name = 'Arial'
run.font.size = Pt(9)
run.font.bold = True
#header_性別
run = head_tables[1].cell(1,5).paragraphs[0].add_run(str(patient_report_ID))
run.font.name = 'Arial'
run.font.size = Pt(9)
run.font.bold = True
#header_報告編號
run = head_tables[1].cell(2,3).paragraphs[0].add_run(str(patient_hospital))
#run.font.name = 'Microsoft JhengHei'
run.font.size = Pt(9)
run.font.bold = True
#header_委託機構
run = doc.tables[0].cell(0,4).paragraphs[0].add_run(str(patient_report_ID))
run.font.name = 'Arial'
run.font.size = Pt(9)
run.font.bold = True
#檢體編號
run = doc.tables[0].cell(2,2).paragraphs[0].add_run(str(patient_name))
#run.font.name = 'Microsoft JhengHei'
#font = '微軟正黑體'
run.font.size = Pt(9)
run.font.bold = True
#doc.tables[0].cell(2, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
#姓名
run = doc.tables[0].cell(2,4).paragraphs[0].add_run(str(patient_birthday))
run.font.name = 'Arial'
run.font.size = Pt(9)
run.font.bold = True
#生日
run = doc.tables[0].cell(3,2).paragraphs[0].add_run(str(patient_sex))
run.font.name = 'Microsoft JhengHei'
run.font.name = 'Arial'
run.font.size = Pt(9)
run.font.bold = True
#性別
run = doc.tables[0].cell(3,4).paragraphs[0].add_run(str(patient_ID))
run.font.name = 'Arial'
run.font.size = Pt(9)
run.font.bold = True
#身分證字號
run = doc.tables[0].cell(4,2).paragraphs[0].add_run(str(patient_sample_type))
#run.font.name = 'Microsoft JhengHei'
run.font.name = 'Arial'
run.font.size = Pt(9)
run.font.bold = True
#檢體類型
run = doc.tables[0].cell(4,4).paragraphs[0].add_run(str(patient_hospital))
#font = '微軟正黑體'
#run.font.name = 'Microsoft JhengHei'
run.font.size = Pt(9)
run.font.bold = True
#委託機構
run = doc.tables[0].cell(5,2).paragraphs[0].add_run(str(patient_number))
run.font.name = 'Arial'
run.font.size = Pt(9)
run.font.bold = True
#病歷編號
run = doc.tables[0].cell(5,4).paragraphs[0].add_run(str(patient_doctor))
run.font.name = 'Microsoft JhengHei'
run.font.size = Pt(9)
run.font.bold = True
#醫生
run = doc.tables[0].cell(6,2).paragraphs[0].add_run(str(patient_test_time))
run.font.name = 'Arial'
run.font.size = Pt(9)
run.font.bold = True
#採集日期
run = doc.tables[0].cell(6,4).paragraphs[0].add_run(str(patient_contact))
#run.font.name = 'Microsoft JhengHei'
run.font.name = 'Arial'
run.font.size = Pt(9)
run.font.bold = True
#聯絡人
run = doc.tables[0].cell(7,2).paragraphs[0].add_run(str(patient_receive_time))
run.font.name = 'Arial'
run.font.size = Pt(9)
run.font.bold = True
#收檢日期
run = doc.tables[0].cell(7,4).paragraphs[0].add_run(str(doctor_phone))
run.font.name = 'Arial'
run.font.size = Pt(9)
run.font.bold = True
#電話
run = doc.tables[0].cell(8,2).paragraphs[0].add_run(str(patient_report_time))
run.font.name = 'Arial'
run.font.size = Pt(9)
run.font.bold = True
#報告日期
run = doc.tables[0].cell(8,4).paragraphs[0].add_run(str(doctor_email))
run.font.name = 'Arial'
run.font.size = Pt(9)
run.font.bold = True
#信箱
doc.save('/home/jimmy/report/patient_info/PFI_done.docx')
